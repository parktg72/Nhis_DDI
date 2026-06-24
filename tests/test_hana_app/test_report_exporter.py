"""tests for hana_app.core.report_exporter"""
import io
import pytest
import pandas as pd

from hana_app.core.report_exporter import (
    build_csv_bytes,
    build_docx_bytes,
    DOCX_AVAILABLE,
    MPL_AVAILABLE,
    _collect_page4_docx_sections,
    _derive_reason,
    _effective_label,
    _summarize_misclassification_reasons,
    _yellow_summary,
)


def _row(**kw):
    defaults = {
        "patient_id": "P001",
        "risk_level": "Yellow",
        "yellow_subtype": "",
        "drug_count": 5,
        "ddi_contraindicated": 0,
        "ddi_major": 0,
        "ddi_moderate": 0,
        "ddi_minor": 0,
        "dup_same_ingredient": 0,
        "institution_count": 1,
        "has_high_risk_drug": 0,
        "has_renal_risk_drug": 0,
        "has_hepatic_risk_drug": 0,
        "triple_whammy": 0,
    }
    return {**defaults, **kw}


# ── CSV ─────────────────────────────────────────────────────────────────────

def test_csv_bytes_red_only():
    df = pd.DataFrame([
        _row(patient_id="P001", risk_level="Red", ddi_contraindicated=2),
        _row(patient_id="P002", risk_level="Green"),
    ])
    b = build_csv_bytes(df)
    assert b[:3] == b"\xef\xbb\xbf", "must start with UTF-8 BOM"
    text = b.decode("utf-8-sig")
    assert "P001" in text
    assert "P002" not in text
    assert "금기 DDI 2건" in text  # _derive_reason: RED_CONTRAINDICATED 건수 포함
    assert "즉각 개입" in text


def test_csv_bytes_all_labels():
    df = pd.DataFrame([
        _row(patient_id="R1", risk_level="Red", ddi_contraindicated=1),
        _row(patient_id="M1", risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR", ddi_major=3),
        _row(patient_id="T1", risk_level="Yellow", yellow_subtype="Y_TRIPLE", drug_count=11),
        _row(patient_id="G1", risk_level="Green"),
    ])
    b = build_csv_bytes(df)
    text = b.decode("utf-8-sig")
    assert "R1" in text and "M1" in text and "T1" in text
    assert "G1" not in text
    assert "즉각 개입" in text
    assert "약사 전화" in text
    assert "문자 안내" in text
    assert "금기 DDI 1건" in text   # RED_CONTRAINDICATED 건수
    assert "중증 DDI 3건" in text   # DDI_MAJOR 건수


def test_csv_bytes_no_target_rows():
    df = pd.DataFrame([_row(patient_id="G1", risk_level="Green")])
    b = build_csv_bytes(df)
    text = b.decode("utf-8-sig")
    assert "G1" not in text
    # header only (환자ID column present)
    assert "환자ID" in text


def test_csv_utf8_bom():
    df = pd.DataFrame([_row(risk_level="Red")])
    b = build_csv_bytes(df)
    assert b[:3] == b"\xef\xbb\xbf"


def test_derive_reason_red():
    r = _row(risk_level="Red", ddi_contraindicated=3)
    assert _derive_reason(r) == "금기 DDI 3건"


def test_derive_reason_major():
    # DDI_MAJOR trigger: ddi_major >= 1
    r = _row(risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR", ddi_major=5)
    assert _derive_reason(r) == "중증 DDI 5건"


def test_derive_reason_triple_with_trigger():
    # SEV_10DRUG_HIGHRISK: drug_count >= 10 AND has_high_risk_drug
    r = _row(risk_level="Yellow", yellow_subtype="Y_TRIPLE",
             drug_count=12, has_high_risk_drug=1)
    reason = _derive_reason(r)
    assert "10종↑+고위험약" in reason


def test_derive_reason_no_trigger_green_path():
    # drug_count >= 5, no other triggers → Green 사유 문구
    r = _row(risk_level="Green", drug_count=12)
    reason = _derive_reason(r)
    assert "5종↑" in reason


# ── DOCX ────────────────────────────────────────────────────────────────────

@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_bytes_returns_bytes():
    last_result = {"model_name": "test_model", "metrics": {"f1_macro": 0.85}}
    b = build_docx_bytes(last_result)
    assert isinstance(b, bytes)
    assert b[:4] == b"PK\x03\x04", "must be a ZIP (DOCX magic)"


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_missing_ddi_means():
    last_result = {"model_name": "m", "metrics": {}}
    b = build_docx_bytes(last_result)
    assert b[:4] == b"PK\x03\x04"


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_empty_feature_importance():
    last_result = {"model_name": "m", "metrics": {}, "feature_importance": []}
    b = build_docx_bytes(last_result)
    assert b[:4] == b"PK\x03\x04"


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_with_features_df():
    df = pd.DataFrame([
        _row(patient_id="R1", risk_level="Red"),
        _row(patient_id="Y1", risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR"),
        _row(patient_id="G1", risk_level="Green"),
    ])
    last_result = {
        "model_name": "hier",
        "metrics": {"f1_macro": 0.9},
        "ddi_means": {"ddi_major": 1.2, "ddi_contraindicated": 0.1},
        "drug_count_stats": {"mean": 6.5, "max": 14},
    }
    b = build_docx_bytes(last_result, df)
    assert b[:4] == b"PK\x03\x04"


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_includes_page4_analysis_sections_and_charts():
    from docx import Document

    df = pd.DataFrame([
        _row(patient_id="R1", risk_level="Red", yellow_subtype="", drug_count=10,
             age=80, sex_m=1, ddi_contraindicated=1, ddi_major=2),
        _row(patient_id="M1", risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR",
             drug_count=8, age=72, sex_m=0, ddi_major=1),
        _row(patient_id="T1", risk_level="Yellow", yellow_subtype="Y_TRIPLE",
             drug_count=12, age=77, sex_m=1, triple_whammy=1),
        _row(patient_id="G1", risk_level="Green", yellow_subtype="", drug_count=5,
             age=66, sex_m=0),
    ])
    last_result = {
        "model_name": "rich_model",
        "target": "risk_binary",
        "metrics": {
            "accuracy": 0.81,
            "f1_macro": 0.82,
            "roc_auc": 0.83,
            "cv_mean": 0.80,
            "cv_std": 0.02,
            "train_size": 100,
            "test_size": 40,
            "confusion_matrix": [[20, 3], [4, 13]],
            "classes": [0, 1],
            "cv_scores": [0.78, 0.80, 0.82],
            "roc_curve": {"fpr": [0.0, 0.2, 1.0], "tpr": [0.0, 0.75, 1.0]},
            "classification_report": "precision recall f1-score\n위험 0.80 0.76 0.78",
        },
        "feature_importance": [
            {"feature": "drug_count", "importance": 0.4},
            {"feature": "ddi_major", "importance": 0.3},
        ],
        "ddi_means": {"ddi_contraindicated": 0.25, "ddi_major": 0.75, "ddi_moderate": 0.0, "ddi_minor": 0.0},
        "drug_count_stats": {"mean": 8.75, "max": 12},
    }
    saved_results = [
        {"timestamp": "20260618_010000", "model_name": "rich_model", "target": "risk_binary",
         "metrics": {"accuracy": 0.81, "f1_macro": 0.82, "roc_auc": 0.83, "cv_mean": 0.80, "train_size": 100}},
        {"timestamp": "20260618_020000", "model_name": "baseline", "target": "risk_binary",
         "metrics": {"accuracy": 0.70, "f1_macro": 0.71, "roc_auc": 0.72, "cv_mean": 0.69, "train_size": 100}},
    ]

    b = build_docx_bytes(last_result, df, saved_results=saved_results)
    doc = Document(io.BytesIO(b))
    body = "\n".join(p.text for p in doc.paragraphs)

    for heading in [
        "혼동 행렬",
        "교차검증",
        "ROC Curve",
        "분류 보고서",
        "모델 비교",
        "분석 대상 정보",
    ]:
        assert heading in body
    assert "precision recall f1-score" in body
    assert "학습 100건" in body

    if MPL_AVAILABLE:
        image_rels = [r for r in doc.part.rels.values() if "image" in r.reltype]
        assert len(image_rels) >= 9


def test_page4_docx_section_plan_covers_all_analysis_tabs():
    df = pd.DataFrame([
        _row(risk_level="Red", yellow_subtype="", drug_count=10, age=80, ddi_contraindicated=1),
        _row(risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR", drug_count=8, age=72, ddi_major=1),
    ])
    result = {
        "metrics": {
            "confusion_matrix": [[9, 1], [2, 8]],
            "cv_scores": [0.78, 0.82],
            "roc_curve": {"fpr": [0, 1], "tpr": [0, 1]},
            "classification_report": "precision recall f1-score",
        },
        "feature_importance": [{"feature": "drug_count", "importance": 0.4}],
        "ddi_means": {"ddi_major": 0.5},
    }
    saved = [{"model_name": "m1", "metrics": {"accuracy": 0.8}}, {"model_name": "m2", "metrics": {"accuracy": 0.7}}]

    sections = _collect_page4_docx_sections(result, df, saved)
    ids = {s["id"] for s in sections}

    assert {
        "feature_importance",
        "confusion_matrix",
        "cross_validation",
        "roc_curve",
        "risk_distribution",
        "risk_count_bar",
        "drug_count_distribution",
        "ddi_severity",
        "yellow_subtype",
        "classification_report",
        "model_comparison",
        "analysis_subject",
    } <= ids


def test_build_docx_bytes_accepts_saved_results_for_model_comparison():
    import inspect

    sig = inspect.signature(build_docx_bytes)

    assert "saved_results" in sig.parameters


def test_build_docx_bytes_accepts_training_results_for_sequential_report():
    import inspect

    sig = inspect.signature(build_docx_bytes)

    assert "training_results" in sig.parameters


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_records_current_sequential_training_results():
    from docx import Document

    last_result = {"model_name": "lightgbm", "metrics": {"f1_macro": 0.82}}
    training_results = {
        "xgboost": {
            "model_name": "xgboost",
            "target": "risk_label",
            "metrics": {"accuracy": 0.81, "f1_macro": 0.80, "roc_auc_ovr": 0.83, "cv_mean": 0.79, "train_size": 120},
        },
        "lightgbm": {
            "model_name": "lightgbm",
            "target": "risk_label",
            "metrics": {"accuracy": 0.84, "f1_macro": 0.82, "roc_auc_ovr": 0.85, "cv_mean": 0.81, "train_size": 120},
        },
    }

    b = build_docx_bytes(last_result, training_results=training_results)
    doc = Document(io.BytesIO(b))
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )

    assert "이번 순차 학습 결과" in text
    assert "xgboost" in text
    assert "lightgbm" in text
    assert "0.8200" in text


def _docx_block_texts(doc):
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, doc).text
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            yield "table", "\n".join(cell.text for row in table.rows for cell in row.cells)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_docx_model_comparison_uses_current_ml_and_dl_training_results():
    from docx import Document

    last_result = {"model_name": "gnn", "target": "risk_binary", "metrics": {"f1_macro": 0.79}}
    training_results = {
        "xgboost": {
            "model_name": "xgboost",
            "target": "risk_binary",
            "metrics": {"accuracy": 0.91, "f1_macro": 0.90, "roc_auc": 0.93, "cv_mean": 0.89, "train_size": 120},
        },
        "lightgbm": {
            "model_name": "lightgbm",
            "target": "risk_binary",
            "metrics": {"accuracy": 0.92, "f1_macro": 0.91, "roc_auc": 0.94, "cv_mean": 0.90, "train_size": 120},
        },
        "tabnet": {
            "model_name": "tabnet",
            "target": "risk_binary",
            "metrics": {"accuracy": 0.88, "f1_macro": 0.86, "roc_auc": 0.89, "cv_mean": 0.85, "train_size": 120},
        },
        "gnn": {
            "model_name": "gnn",
            "target": "risk_binary",
            "metrics": {"accuracy": 0.87, "f1_macro": 0.79, "roc_auc": 0.88, "cv_mean": 0.84, "train_size": 120},
        },
    }

    b = build_docx_bytes(last_result, saved_results=[], training_results=training_results)
    doc = Document(io.BytesIO(b))
    blocks = list(_docx_block_texts(doc))
    heading_idx = next(i for i, (_, text) in enumerate(blocks) if text == "5-5. 모델 비교")
    comparison_table = next(text for kind, text in blocks[heading_idx + 1:] if kind == "table")

    assert "Accuracy" in comparison_table
    assert "F1" in comparison_table
    assert "AUC" in comparison_table
    for model_name in ["xgboost", "lightgbm", "tabnet", "gnn"]:
        assert model_name in comparison_table
    if MPL_AVAILABLE:
        image_rels = [r for r in doc.part.rels.values() if "image" in r.reltype]
        assert image_rels, "Accuracy/F1/AUC model-comparison chart should be embedded"


def test_page4_docx_section_plan_treats_current_training_results_as_model_comparison():
    result = {"metrics": {"f1_macro": 0.8}}
    training_results = {
        "xgboost": {"model_name": "xgboost", "metrics": {"accuracy": 0.9}},
        "tabnet": {"model_name": "tabnet", "metrics": {"accuracy": 0.8}},
    }

    sections = _collect_page4_docx_sections(
        result,
        pd.DataFrame([_row(risk_level="Red")]),
        saved_results=[],
        training_results=training_results,
    )

    assert "model_comparison" in {s["id"] for s in sections}


def test_docx_yellow_action_summary_matches_page4_for_y_other():
    df = pd.DataFrame([
        _row(risk_level="Yellow", yellow_subtype="Y_OTHER"),
        _row(risk_level="Yellow", yellow_subtype="Y_DDI_MAJOR"),
    ])

    summary = _yellow_summary(df)
    action_by_label = dict(zip(summary["yellow_subtype"], summary["action"]))

    assert action_by_label["Y_OTHER"] == "알림 없음"
    assert action_by_label["Y_DDI_MAJOR"] == "약사 전화"


def test_misclassification_reason_summary_redacts_identifiers():
    df = pd.DataFrame([
        _row(patient_id="SECRET001", risk_level="Red", drug_count=12, ddi_major=2, red_suspect=True),
        _row(patient_id="SECRET002", risk_level="Green", drug_count=4, ddi_major=0),
        _row(patient_id="SECRET003", risk_level="Yellow", drug_count=9, ddi_moderate=3),
    ])
    metrics = {
        "misclassified_cases": [
            {"row_index": 0, "actual": "Red", "predicted": "Yellow", "probability": 0.48},
            {"patient_id": "SECRET002", "actual": "Green", "predicted": "Red", "probability": 0.91},
            {"row_index": 2, "actual": "Yellow", "predicted": "Green", "probability": 0.52},
        ]
    }

    summary = _summarize_misclassification_reasons(metrics, df)

    assert summary["total"] == 3
    assert summary["type_counts"]["FN 고위험 과소예측"] == 1
    assert summary["type_counts"]["FP 위험 과대예측"] == 1
    assert summary["type_counts"]["인접 위험단계 오분류"] == 1
    assert len(summary["examples"]) == 3
    joined = "\n".join(str(x) for x in summary["examples"])
    assert "SECRET" not in joined
    assert "patient_id" not in joined
    assert "약물 수" in joined
    assert "중증 DDI" in joined


def test_docx_section_plan_includes_misclassification_analysis():
    result = {
        "metrics": {
            "confusion_matrix": [[8, 2], [3, 7]],
            "misclassified_cases": [{"row_index": 0, "actual": "Red", "predicted": "Yellow"}],
        }
    }
    sections = _collect_page4_docx_sections(result, pd.DataFrame([_row(risk_level="Red")]), None)
    assert "misclassification_analysis" in {s["id"] for s in sections}
