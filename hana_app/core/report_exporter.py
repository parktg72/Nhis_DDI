"""
보고서 내보내기 — 종합 서비스 DOCX 보고서 + 대상자 CSV.

공개 API:
    build_csv_bytes(features_df)              -> bytes  (UTF-8 BOM)
    build_docx_bytes(last_result, features_df) -> bytes  (차트 포함 종합 보고서)
    DOCX_AVAILABLE                            -> bool
    MPL_AVAILABLE                             -> bool
"""
from __future__ import annotations

import datetime as dt
import io
from types import SimpleNamespace
from typing import Optional

import pandas as pd

# ── matplotlib (차트 생성) ────────────────────────────────────────────────────
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    matplotlib.rcParams["font.family"] = "Malgun Gothic"
    matplotlib.rcParams["axes.unicode_minus"] = False
    MPL_AVAILABLE = True
except ImportError:
    MPL_AVAILABLE = False

# ── python-docx ──────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from scripts.etl.clinical_rules import (
        collect_red_triggers as _collect_red,
    )
    from scripts.etl.clinical_rules import (
        collect_severe_immediate_triggers as _collect_severe,
    )
    from scripts.etl.clinical_rules import (
        collect_yellow_triggers as _collect_yellow,
    )
    _CLINICAL_RULES_AVAILABLE = True
except ImportError:
    _CLINICAL_RULES_AVAILABLE = False

_summarize_yellow_subtypes = None
_summarize_actions = None
try:
    from hana_app.core.yellow_subtype_view import (
        summarize_actions as _summarize_actions,
    )
    from hana_app.core.yellow_subtype_view import (
        summarize_yellow_subtypes as _summarize_yellow_subtypes,
    )
    _YELLOW_VIEW_AVAILABLE = True
except ImportError:
    _YELLOW_VIEW_AVAILABLE = False

# ── 색상 팔레트 ──────────────────────────────────────────────────────────────
_C = {
    "Red":       "#e74c3c",
    "MAJOR":     "#e67e22",
    "TRIPLE":    "#f1c40f",
    "monitor":   "#3498db",
    "noalert":   "#bdc3c7",
    "Green":     "#27ae60",
    "Yellow":    "#f39c12",
    "Normal":    "#95a5a6",
    "fi_bar":    "#2980b9",
    "ddi_contra":"#e74c3c",
    "ddi_major": "#e67e22",
    "ddi_mod":   "#f1c40f",
    "ddi_minor": "#3498db",
}

_INTERVENTION_KO = {
    "Red":       "즉각 개입",
    "Y_DDI_MAJOR": "약사 전화",
    "Y_TRIPLE":  "문자 안내",
}

_INTERVENTION_ROWS = [
    ("즉각 개입",  "Red",                        _C["Red"]),
    ("약사 전화",  "Y_DDI_MAJOR",                _C["MAJOR"]),
    ("문자 안내",  "Y_TRIPLE",                   _C["TRIPLE"]),
    ("모니터링",   "Y_DOUBLE·Y_DDI_MOD·Y_DUP·Y_FRAG", _C["monitor"]),
    ("관여안함",   "No_Alert·Green·Normal",       _C["noalert"]),  # _count_distribution 키와 일치 필수
]

_MONITORING_SUBTYPES = frozenset({"Y_DOUBLE", "Y_DDI_MOD", "Y_DUP", "Y_FRAG"})

_REASON_TOKEN_KO = {
    "RED_CONTRAINDICATED": "금기 DDI",
    "SEV_TRIPLE_WHAMMY":   "Triple Whammy",
    "SEV_10DRUG_HIGHRISK": "10종↑+고위험약",
    "SEV_ELDERLY_ORGAN":   "고령(75+)+장기부전",
    "DDI_MAJOR":           "중증 DDI",
    "DDI_MOD":             "Moderate DDI",
    "DUP":                 "동일성분 중복",
    "FRAG":                "다기관(3곳↑)",
}

_FEAT_LABELS = {
    "drug_count": "총 약물 수",
    "drug_count_7d": "최근 7일 복용",
    "institution_count": "처방 기관 수",
    "ddi_contraindicated": "금기 DDI",
    "ddi_major": "Major DDI",
    "ddi_moderate": "Moderate DDI",
    "ddi_minor": "Minor DDI",
    "triple_whammy": "Triple Whammy",
    "qt_risk_count": "QT 위험 약물",
    "dup_same_ingredient": "동일성분 중복",
    "dup_atc5": "ATC5 중복",
    "dup_atc4": "ATC4 중복",
    "dup_atc3": "ATC3 중복",
    "dup_efmdc": "약효분류 중복",
    "age": "연령",
    "sex_m": "성별(남)",
}


# ── 내부 헬퍼 ────────────────────────────────────────────────────────────────

def _fig_to_png(fig) -> bytes:
    canvas = FigureCanvasAgg(fig)
    buf = io.BytesIO()
    canvas.print_png(buf)
    plt.close(fig)
    return buf.getvalue()


def _count_distribution(features_df: pd.DataFrame) -> dict[str, int]:
    ys = (features_df["yellow_subtype"]
          if "yellow_subtype" in features_df.columns
          else pd.Series("", index=features_df.index))
    red_n  = int((features_df["risk_level"] == "Red").sum())
    maj_n  = int((ys == "Y_DDI_MAJOR").sum())
    tri_n  = int((ys == "Y_TRIPLE").sum())
    mon_n  = int(ys.isin(_MONITORING_SUBTYPES).sum())
    none_n = max(0, len(features_df) - red_n - maj_n - tri_n - mon_n)
    return {"Red": red_n, "Y_DDI_MAJOR": maj_n, "Y_TRIPLE": tri_n,
            "모니터링": mon_n, "관여안함": none_n}


def _chart_intervention(counts: dict, total: int) -> bytes:
    """개입 위계 수평 바차트."""
    labels   = ["즉각 개입\n(Red)", "약사 전화\n(Y_DDI_MAJOR)", "문자 안내\n(Y_TRIPLE)",
                "모니터링", "관여 안함"]
    keys     = ["Red", "Y_DDI_MAJOR", "Y_TRIPLE", "모니터링", "관여안함"]
    colors   = [_C["Red"], _C["MAJOR"], _C["TRIPLE"], _C["monitor"], _C["noalert"]]
    values   = [counts.get(k, 0) for k in keys]
    pcts     = [v / total * 100 if total else 0 for v in values]

    fig, ax = plt.subplots(figsize=(7, 3.2))
    bars = ax.barh(labels[::-1], values[::-1], color=colors[::-1], edgecolor="white")
    for bar, pct, val in zip(bars, pcts[::-1], values[::-1]):
        ax.text(bar.get_width() + max(values) * 0.01, bar.get_y() + bar.get_height() / 2,
                f"{val:,}명 ({pct:.1f}%)", va="center", fontsize=9)
    ax.set_xlabel("환자 수")
    ax.set_title("개입 위계별 환자 분포")
    ax.set_xlim(0, (max(values) * 1.25) if max(values, default=0) > 0 else 1)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_risk_pie(features_df: pd.DataFrame) -> bytes:
    """위험도 분포 파이차트."""
    dist = features_df["risk_level"].value_counts()
    labels = list(dist.index)
    values = list(dist.values)
    colors = [_C.get(l, "#aaaaaa") for l in labels]

    fig, ax = plt.subplots(figsize=(5, 4))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors,
        autopct="%1.1f%%", startangle=140,
        pctdistance=0.75, wedgeprops={"edgecolor": "white", "linewidth": 1.5}
    )
    for t in autotexts:
        t.set_fontsize(9)
    ax.set_title("위험도 분포")
    fig.tight_layout()
    return _fig_to_png(fig)





def _chart_risk_bar_from_counts(counts: dict[str, int]) -> bytes:
    """위험도/개입 카운트 막대차트."""
    items = [(str(k), int(v)) for k, v in counts.items() if int(v) >= 0]
    if not items:
        return b""
    labels, values = zip(*items)
    colors = [_C.get(l, _C.get("monitor", "#3498db")) for l in labels]
    fig, ax = plt.subplots(figsize=(6, 3.2))
    bars = ax.bar(labels, values, color=colors, edgecolor="white")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.02,
                f"{val:,}", ha="center", va="bottom", fontsize=9)
    ax.set_ylabel("환자 수")
    ax.set_title("위험도별 환자 수")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)





def _chart_pie_from_counts(counts: dict[str, int], title: str) -> bytes:
    items = [(str(k), int(v)) for k, v in counts.items() if int(v) > 0]
    if not items:
        return b""
    labels, values = zip(*items)
    colors = [_C.get(l, _C.get("monitor", "#3498db")) for l in labels]
    fig, ax = plt.subplots(figsize=(5, 4))
    ax.pie(values, labels=labels, colors=colors, autopct="%1.1f%%",
           startangle=140, wedgeprops={"edgecolor": "white", "linewidth": 1.5})
    ax.set_title(title)
    fig.tight_layout()
    return _fig_to_png(fig)


def _yellow_summary(features_df: pd.DataFrame) -> pd.DataFrame:
    if _YELLOW_VIEW_AVAILABLE and _summarize_yellow_subtypes is not None:
        return _summarize_yellow_subtypes(features_df)
    if "yellow_subtype" not in features_df.columns:
        return pd.DataFrame(columns=["yellow_subtype", "count", "action"])
    s = features_df["yellow_subtype"].dropna()
    s = s[s.astype(str) != ""]
    if s.empty:
        return pd.DataFrame(columns=["yellow_subtype", "count", "action"])
    counts = s.value_counts().reset_index()
    counts.columns = ["yellow_subtype", "count"]
    action_map = {
        "Y_DDI_MAJOR": "약사 전화",
        "Y_TRIPLE": "문자 안내",
        "Y_DOUBLE": "문자 알림",
        "Y_DDI_MOD": "문자 알림",
        "Y_DUP": "문서 + 문자 알림",
        "Y_FRAG": "문자 알림",
    }
    counts["action"] = counts["yellow_subtype"].map(lambda x: action_map.get(str(x), "알림 없음"))
    return counts


def _chart_yellow_pie(features_df: pd.DataFrame) -> bytes:
    summary = _yellow_summary(features_df)
    if summary.empty:
        return b""
    return _chart_pie_from_counts(dict(zip(summary["yellow_subtype"], summary["count"])), "Yellow 세부 라벨 분포")


def _chart_action_distribution(features_df: pd.DataFrame) -> bytes:
    if _YELLOW_VIEW_AVAILABLE and _summarize_actions is not None:
        action_df = _summarize_actions(features_df)
        counts = dict(zip(action_df.get("action", []), action_df.get("count", [])))
        return _chart_risk_bar_from_counts(counts)
    frames = []
    summary = _yellow_summary(features_df)
    if not summary.empty:
        frames.append(summary[["action", "count"]])
    if "risk_level" in features_df.columns:
        red_n = int((features_df["risk_level"] == "Red").sum())
        if red_n:
            frames.append(pd.DataFrame({"action": ["즉각 개입"], "count": [red_n]}))
    if not frames:
        return b""
    counts = pd.concat(frames, ignore_index=True).groupby("action")["count"].sum().to_dict()
    return _chart_risk_bar_from_counts(counts)


def _chart_ddi(ddi_means: dict) -> bytes:
    """DDI 심각도별 평균 바차트."""
    label_map = [
        ("ddi_contraindicated", "금기",    _C["ddi_contra"]),
        ("ddi_major",           "Major",   _C["ddi_major"]),
        ("ddi_moderate",        "Moderate",_C["ddi_mod"]),
        ("ddi_minor",           "Minor",   _C["ddi_minor"]),
    ]
    items = [(lbl, col, ddi_means[k]) for k, lbl, col in label_map if k in ddi_means]
    if not items:
        return b""
    labels, colors, vals = zip(*items)

    fig, ax = plt.subplots(figsize=(5.5, 3))
    bars = ax.bar(labels, vals, color=colors, edgecolor="white", width=0.5)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(vals) * 0.02,
                f"{v:.2f}", ha="center", va="bottom", fontsize=9)
    ax.set_ylabel("평균 쌍 수")
    ax.set_title("DDI 심각도별 평균 쌍 수")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_feature_importance(fi_df: pd.DataFrame, top_n: int = 15) -> bytes:
    """피처 중요도 수평 바차트."""
    fi_top = fi_df.sort_values("importance", ascending=False).head(top_n)
    fi_top = fi_top.sort_values("importance", ascending=True)
    labels = [_FEAT_LABELS.get(str(f), str(f)) for f in fi_top["feature"]]
    values = list(fi_top["importance"])

    fig, ax = plt.subplots(figsize=(7, max(3, len(labels) * 0.38)))
    ax.barh(labels, values, color=_C["fi_bar"], edgecolor="white")
    ax.set_xlabel("중요도")
    ax.set_title(f"피처 중요도 Top {top_n}")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)





def _chart_confusion_matrix(cm, classes=None) -> bytes:
    arr = pd.DataFrame(cm).to_numpy()
    if arr.size == 0:
        return b""
    labels = [str(c) for c in (classes or range(arr.shape[0]))]
    fig, ax = plt.subplots(figsize=(4.8, 4.0))
    im = ax.imshow(arr, cmap="Blues")
    ax.set_xticks(range(len(labels)), labels=labels)
    ax.set_yticks(range(len(labels)), labels=labels)
    ax.set_xlabel("예측")
    ax.set_ylabel("실제")
    ax.set_title("혼동 행렬")
    for i in range(arr.shape[0]):
        for j in range(arr.shape[1]):
            ax.text(j, i, str(arr[i, j]), ha="center", va="center", fontsize=9)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_cv_scores(cv_scores: list[float]) -> bytes:
    if not cv_scores:
        return b""
    labels = [f"Fold {i + 1}" for i in range(len(cv_scores))]
    mean_v = sum(cv_scores) / len(cv_scores)
    fig, ax = plt.subplots(figsize=(6, 3.2))
    bars = ax.bar(labels, cv_scores, color=_C["fi_bar"], edgecolor="white")
    ax.axhline(mean_v, color=_C["Red"], linestyle="--", linewidth=1.2, label=f"평균 {mean_v:.4f}")
    for bar, val in zip(bars, cv_scores):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.005,
                f"{val:.3f}", ha="center", va="bottom", fontsize=8)
    ax.set_ylabel("Score")
    ax.set_title("교차검증 결과")
    ax.legend(fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_roc_curve(roc_data: dict, auc_value=None) -> bytes:
    if not roc_data or "fpr" not in roc_data or "tpr" not in roc_data:
        return b""
    fig, ax = plt.subplots(figsize=(5, 4))
    label = "ROC" + (f" (AUC={auc_value:.4f})" if isinstance(auc_value, (int, float)) else "")
    ax.plot(roc_data["fpr"], roc_data["tpr"], color=_C["fi_bar"], linewidth=2, label=label)
    ax.plot([0, 1], [0, 1], color="#7f8c8d", linestyle="--", linewidth=1, label="Random")
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.set_title("ROC Curve")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_model_comparison(saved_results: list[dict]) -> bytes:
    rows = []
    for r in saved_results or []:
        m = r.get("metrics", {}) or {}
        model_label = str(r.get("model_name", "?"))
        family = r.get("model_family")
        if family:
            model_label = f"{model_label} ({family})"
        rows.append({
            "model": model_label,
            "Accuracy": _metric_float(m.get("accuracy")),
            "F1": _metric_float(m.get("f1_macro")),
            "AUC": _metric_float(m.get("roc_auc", m.get("roc_auc_ovr"))),
        })
    if len(rows) < 2:
        return b""
    df = pd.DataFrame(rows)
    fig, ax = plt.subplots(figsize=(7, 3.5))
    x = range(len(df))
    width = 0.25
    for idx, col in enumerate(["Accuracy", "F1", "AUC"]):
        ax.bar([v + (idx - 1) * width for v in x], df[col], width=width, label=col)
    ax.set_xticks(list(x), df["model"], rotation=20, ha="right")
    ax.set_ylim(0, 1)
    ax.set_ylabel("Score")
    ax.set_title("모델별 성능 비교")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return _fig_to_png(fig)




_RISK_RANK = {"Normal": 0, "정상": 0, "Green": 1, "Yellow": 2, "Red": 3, "위험": 3, "위험군": 3, 0: 0, 1: 3, 2: 2, 3: 3}
_SAFE_MISCLS_FEATURES = [
    "drug_count", "ddi_contraindicated", "ddi_major", "ddi_moderate", "ddi_minor",
    "dup_same_ingredient", "institution_count", "has_high_risk_drug",
    "has_renal_risk_drug", "has_hepatic_risk_drug", "triple_whammy",
    "yellow_subtype", "red_suspect",
]


def _risk_rank(label) -> int:
    return _RISK_RANK.get(label, _RISK_RANK.get(str(label), 0))


def _misclassification_type(actual, predicted) -> str:
    a = _risk_rank(actual)
    p = _risk_rank(predicted)
    if str(actual) == "Red" and p < a:
        return "FN 고위험 과소예측"
    if str(predicted) == "Red" and p > a:
        return "FP 위험 과대예측"
    if a > p and a >= 2 and abs(a - p) > 1:
        return "FN 고위험 과소예측"
    if p > a and p >= 2 and abs(p - a) > 1:
        return "FP 위험 과대예측"
    return "인접 위험단계 오분류"


def _safe_feature_snapshot(row: dict | pd.Series | None) -> dict:
    if row is None:
        return {}
    data = row.to_dict() if hasattr(row, "to_dict") else dict(row)
    return {k: data.get(k) for k in _SAFE_MISCLS_FEATURES if k in data and pd.notna(data.get(k))}


def _reason_from_safe_features(features: dict) -> str:
    parts: list[str] = []
    try:
        dc = int(features.get("drug_count", 0) or 0)
        if dc >= 10:
            parts.append(f"약물 수 {dc}종(10종 이상)")
        elif dc >= 5:
            parts.append(f"약물 수 {dc}종(5종 이상)")
    except Exception:
        pass
    for key, label in [
        ("ddi_contraindicated", "금기 DDI"),
        ("ddi_major", "중증 DDI"),
        ("ddi_moderate", "Moderate DDI"),
        ("dup_same_ingredient", "동일성분 중복"),
        ("institution_count", "다기관 처방"),
        ("triple_whammy", "Triple Whammy"),
    ]:
        try:
            val = features.get(key, 0) or 0
            if key == "institution_count" and float(val) < 3:
                continue
            if key != "institution_count" and float(val) <= 0:
                continue
            parts.append(f"{label} {int(float(val))}건")
        except Exception:
            continue
    if features.get("has_high_risk_drug"):
        parts.append("고위험약 신호")
    if features.get("has_renal_risk_drug"):
        parts.append("신장주의 약물 신호")
    if features.get("has_hepatic_risk_drug"):
        parts.append("간장주의 약물 신호")
    if features.get("red_suspect"):
        parts.append("Red 의심 경계구간")
    ys = features.get("yellow_subtype")
    if ys:
        parts.append(f"Yellow subtype={ys}")
    return "; ".join(parts[:5]) if parts else "저장된 안전 feature 기준 명확한 단일 원인 없음"


def _summarize_misclassification_reasons(metrics: dict, features_df: Optional[pd.DataFrame] = None,
                                          max_examples: int = 20) -> dict:
    """식별자를 제외한 오판 요약/사유를 만든다.

    입력은 metrics['misclassified_cases']를 우선 사용한다. 각 case에는 actual/predicted,
    probability/confidence, features 또는 row_index(위치 기반)만 허용한다. patient_id 등
    식별자는 출력하지 않는다.
    """
    cases = list((metrics or {}).get("misclassified_cases") or [])
    examples: list[dict[str, str]] = []
    type_counts: dict[str, int] = {}
    for i, case in enumerate(cases[:max_examples], 1):
        actual = case.get("actual", case.get("y_true", "?"))
        predicted = case.get("predicted", case.get("y_pred", "?"))
        mtype = _misclassification_type(actual, predicted)
        type_counts[mtype] = type_counts.get(mtype, 0) + 1
        features = dict(case.get("features") or {})
        if not features and features_df is not None and "row_index" in case:
            try:
                pos = int(case["row_index"])
                if 0 <= pos < len(features_df):
                    features = _safe_feature_snapshot(features_df.iloc[pos])
            except Exception:
                features = {}
        confidence = case.get("confidence", case.get("probability", case.get("proba", "")))
        if isinstance(confidence, float):
            confidence_s = f"{confidence:.3f}"
        else:
            confidence_s = str(confidence) if confidence not in (None, "") else ""
        examples.append({
            "no": f"오판-{i:03d}",
            "actual": str(actual),
            "predicted": str(predicted),
            "type": mtype,
            "confidence": confidence_s,
            "reason": _reason_from_safe_features(features),
        })
    # max_examples 밖의 케이스도 type_counts에 반영
    for case in cases[max_examples:]:
        mtype = _misclassification_type(case.get("actual", case.get("y_true", "?")),
                                        case.get("predicted", case.get("y_pred", "?")))
        type_counts[mtype] = type_counts.get(mtype, 0) + 1
    return {"total": len(cases), "type_counts": type_counts, "examples": examples}


def _add_png(doc, png_bytes: bytes, width_inches: float = 5.5) -> None:
    if png_bytes:
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for row_data in rows:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)


def _deduplicate_features_by_patient_id(features_df: pd.DataFrame) -> pd.DataFrame:
    """Return patient-level rows for report subject summaries.

    Rows with a non-empty ``patient_id`` are counted once, keeping the first
    occurrence. Rows without a usable patient id are left untouched because
    they cannot be safely deduplicated.
    """
    if "patient_id" not in features_df.columns:
        return features_df

    keys = features_df["patient_id"].astype("string").str.strip()
    has_key = keys.notna() & (keys != "")
    duplicate_patient = has_key & keys.where(has_key).duplicated(keep="first")
    if not bool(duplicate_patient.any()):
        return features_df
    return features_df.loc[~duplicate_patient].copy()


def _chart_yellow_subtype(features_df: pd.DataFrame) -> bytes:
    """Yellow 서브타입 분포 바차트."""
    if "yellow_subtype" not in features_df.columns:
        return b""
    ys = features_df["yellow_subtype"].dropna()
    ys = ys[ys.astype(str) != ""]
    if ys.empty:
        return b""
    dist = ys.value_counts()
    _ys_colors = {
        "Y_TRIPLE": _C["TRIPLE"], "Y_DOUBLE": _C["monitor"],
        "Y_DDI_MAJOR": _C["MAJOR"], "Y_DDI_MOD": "#e67e22",
        "Y_DUP": "#9b59b6", "Y_FRAG": "#3498db", "Y_OTHER": _C["noalert"],
    }
    colors = [_ys_colors.get(str(l), "#aaaaaa") for l in dist.index]
    total = dist.sum()
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar([str(l) for l in dist.index], dist.values, color=colors, edgecolor="white")
    for bar, val in zip(bars, dist.values):
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + total * 0.005,
                f"{val:,}\n({val / total * 100:.1f}%)",
                ha="center", va="bottom", fontsize=8)
    ax.set_ylabel("환자 수")
    ax.set_title("Yellow 서브타입 세부 분포")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_drug_hist(features_df: pd.DataFrame) -> bytes:
    """다약제 수 분포 히스토그램."""
    if "drug_count" not in features_df.columns:
        return b""
    vals = features_df["drug_count"].dropna()
    if vals.empty:
        return b""
    n_bins = min(30, max(1, int(vals.max() - vals.min()) + 1))
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.hist(vals, bins=n_bins, color=_C["fi_bar"], edgecolor="white", alpha=0.85)
    ax.axvline(vals.mean(), color=_C["Red"], linestyle="--", linewidth=1.2,
               label=f"평균 {vals.mean():.1f}종")
    ax.set_xlabel("약물 수 (종)")
    ax.set_ylabel("환자 수")
    ax.set_title("다약제 수 분포")
    ax.legend(fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_age_hist(features_df: pd.DataFrame) -> bytes:
    """연령 분포 히스토그램."""
    if "age" not in features_df.columns:
        return b""
    vals = features_df["age"].dropna()
    vals = vals[vals >= 0]
    if vals.empty:
        return b""
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.hist(vals, bins=20, color=_C["Green"], edgecolor="white", alpha=0.85)
    ax.axvline(vals.mean(), color=_C["Red"], linestyle="--", linewidth=1.2,
               label=f"평균 {vals.mean():.1f}세")
    ax.set_xlabel("연령 (세)")
    ax.set_ylabel("환자 수")
    ax.set_title("연령 분포")
    ax.legend(fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _fig_to_png(fig)


# ── CSV ──────────────────────────────────────────────────────────────────────

def _effective_label(row) -> str:
    if row.get("risk_level") == "Red":
        return "Red"
    return row.get("yellow_subtype") or ""


def _derive_reason(row) -> str:
    """clinical_rules 함수로 트리거 재파생 → 한국어 판정사유 (단일 출처 보장)."""
    if not _CLINICAL_RULES_AVAILABLE:
        return ""
    ns = SimpleNamespace(
        ddi_contraindicated=int(row.get("ddi_contraindicated", 0) or 0),
        ddi_major=int(row.get("ddi_major", 0) or 0),
        ddi_moderate=int(row.get("ddi_moderate", 0) or 0),
        dup_same_ingredient=int(row.get("dup_same_ingredient", 0) or 0),
        institution_count=int(row.get("institution_count", 0) or 0),
        triple_whammy=bool(row.get("triple_whammy", 0)),
        drug_count=int(row.get("drug_count", 0) or 0),
        has_high_risk_drug=bool(row.get("has_high_risk_drug", 0)),
        has_renal_risk_drug=bool(row.get("has_renal_risk_drug", 0)),
        has_hepatic_risk_drug=bool(row.get("has_hepatic_risk_drug", 0)),
        age=row.get("age"),
        ddi_minor=int(row.get("ddi_minor", 0) or 0),
    )
    red = _collect_red(ns)
    if red:
        triggers = red
    else:
        triggers = _collect_severe(ns) | _collect_yellow(ns)
    if triggers:
        parts = []
        for t in sorted(triggers):
            ko = _REASON_TOKEN_KO.get(t, t)
            if t == "RED_CONTRAINDICATED":
                parts.append(f"{ko} {ns.ddi_contraindicated}건")
            elif t == "DDI_MAJOR":
                parts.append(f"{ko} {ns.ddi_major}건")
            elif t == "DDI_MOD":
                parts.append(f"{ko} {ns.ddi_moderate}쌍")
            elif t == "DUP":
                parts.append(f"{ko} {ns.dup_same_ingredient}건")
            else:
                parts.append(ko)
        return ", ".join(parts)
    if ns.ddi_minor >= 1:
        return f"Minor DDI {ns.ddi_minor}건"
    if ns.drug_count >= 5:
        return f"5종↑ ({ns.drug_count}종)"
    return ""


def build_csv_bytes(features_df: pd.DataFrame) -> bytes:
    """Red / Y_DDI_MAJOR / Y_TRIPLE 환자 행 → UTF-8 BOM CSV bytes."""
    ys = (features_df["yellow_subtype"]
          if "yellow_subtype" in features_df.columns
          else pd.Series("", index=features_df.index))
    mask = (features_df["risk_level"] == "Red") | ys.isin({"Y_DDI_MAJOR", "Y_TRIPLE"})
    filtered = features_df[mask].copy()

    _PRIORITY = {"Red": 0, "Y_DDI_MAJOR": 1, "Y_TRIPLE": 2}

    if filtered.empty:
        filtered["개입조치"] = pd.Series(dtype=str)
        filtered["위험라벨"] = pd.Series(dtype=str)
        filtered["사유"]    = pd.Series(dtype=str)
    else:
        filtered["개입조치"] = filtered.apply(
            lambda r: _INTERVENTION_KO.get(_effective_label(r), ""), axis=1)
        filtered["위험라벨"] = filtered.apply(_effective_label, axis=1)
        filtered["사유"]    = filtered.apply(_derive_reason, axis=1)
        filtered["_sort"] = filtered["위험라벨"].map(_PRIORITY).fillna(9)
        filtered = filtered.sort_values("_sort").drop(columns=["_sort"])

    out_cols = {
        "patient_id": "환자ID", "개입조치": "개입조치", "위험라벨": "위험라벨",
        "사유": "사유", "drug_count": "다약제수", "ddi_major": "중증DDI건수",
        "ddi_contraindicated": "금기DDI건수", "dup_same_ingredient": "중복처방수",
        "institution_count": "다기관수",
    }
    present = [c for c in out_cols if c in filtered.columns
               or c in ("개입조치", "위험라벨", "사유")]
    buf = io.BytesIO()
    filtered[present].rename(columns=out_cols).to_csv(buf, index=False, encoding="utf-8-sig")
    return buf.getvalue()


# ── DOCX 종합 보고서 ──────────────────────────────────────────────────────────

def _fmt_report_metric(value) -> str:
    if isinstance(value, (int, float)):
        return f"{float(value):.4f}"
    return "—" if value is None else str(value)


def _metric_float(value) -> float:
    try:
        if value is None:
            return 0.0
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _model_family(model_name: str) -> str:
    name = str(model_name).lower()
    if name == "hierarchical":
        return "Hierarchical"
    if name in {"tabnet", "gnn", "temporal_transformer"}:
        return "DL"
    return "ML"


def _iter_training_results(training_results):
    if not training_results:
        return []
    if isinstance(training_results, dict):
        return list(training_results.items())
    return [(str(i + 1), r) for i, r in enumerate(training_results)]


def _result_timestamp(result: dict) -> str:
    return str(result.get("timestamp") or "")


def _comparison_results(saved_results=None, training_results=None) -> list[dict]:
    """DOCX 모델 비교용 결과 병합.

    현재 세션 순차 학습 결과(ML + Phase 3 DL)를 우선 포함하고, 저장 이력은
    모델별 최신 1건만 뒤에 붙인다. 같은 모델의 오래된 이력을 그대로 나열하면
    최근 hierarchical 재학습이 반복되어 ML/DL 비교가 빠진 것처럼 보이므로,
    비교 표는 "시간순 로그"가 아니라 "모델별 최신 성능 비교"로 유지한다.
    """
    merged: list[dict] = []
    seen_models: set[str] = set()

    def _add(result: dict, fallback_key: str = "", fallback_order: int = 0) -> None:
        if not isinstance(result, dict):
            return
        model_name = str(result.get("model_name") or fallback_key or "?")
        model_key = model_name.lower()
        if model_key in seen_models:
            return
        target = str(result.get("target", "?"))
        timestamp = str(result.get("timestamp") or f"current-{fallback_order:02d}")
        seen_models.add(model_key)
        row = dict(result)
        row["model_name"] = model_name
        row["target"] = target
        row["timestamp"] = timestamp
        row["model_family"] = _model_family(model_name)
        merged.append(row)

    for order, (key, result) in enumerate(_iter_training_results(training_results), start=1):
        _add(result, str(key), order)

    latest_saved_by_model: dict[str, dict] = {}
    for result in saved_results or []:
        if not isinstance(result, dict):
            continue
        model_name = str(result.get("model_name") or "?")
        model_key = model_name.lower()
        if model_key in seen_models:
            continue
        prev = latest_saved_by_model.get(model_key)
        if prev is None or _result_timestamp(result) > _result_timestamp(prev):
            latest_saved_by_model[model_key] = result

    latest_saved = sorted(
        latest_saved_by_model.values(),
        key=_result_timestamp,
        reverse=True,
    )
    for order, result in enumerate(latest_saved, start=len(merged) + 1):
        _add(result, str(result.get("model_name", "?")), order)
    return merged


def _training_results_rows(training_results) -> list[list[str]]:
    """현재 세션에서 순차 학습한 모델 결과를 DOCX 표 행으로 정규화."""
    if not training_results:
        return []
    items = _iter_training_results(training_results)
    if len(items) < 2:
        return []

    rows: list[list[str]] = []
    for order, (key, result) in enumerate(items, start=1):
        if not isinstance(result, dict):
            continue
        metrics = result.get("metrics", {}) or {}
        rows.append([
            str(order),
            str(result.get("model_name") or key),
            str(result.get("target", "?")),
            _fmt_report_metric(metrics.get("accuracy")),
            _fmt_report_metric(metrics.get("f1_macro")),
            _fmt_report_metric(metrics.get("roc_auc", metrics.get("roc_auc_ovr"))),
            _fmt_report_metric(metrics.get("cv_mean")),
            str(metrics.get("train_size", "?")),
        ])
    return rows


def _collect_page4_docx_sections(
    last_result: dict,
    features_df: Optional[pd.DataFrame] = None,
    saved_results: Optional[list[dict]] = None,
    training_results: Optional[dict | list[dict]] = None,
) -> list[dict[str, str]]:
    """4_결과_분석 화면의 탭별 DOCX 포함 계획을 반환한다."""
    metrics = last_result.get("metrics", {}) or {}
    sections: list[dict[str, str]] = []
    if last_result.get("feature_importance") is not None:
        sections.append({"id": "feature_importance", "title": "피처 중요도"})
    if metrics.get("confusion_matrix"):
        sections.append({"id": "confusion_matrix", "title": "혼동 행렬"})
    if metrics.get("misclassified_cases"):
        sections.append({"id": "misclassification_analysis", "title": "오판 사유 분석"})
    if metrics.get("cv_scores"):
        sections.append({"id": "cross_validation", "title": "교차검증"})
    if metrics.get("roc_curve"):
        sections.append({"id": "roc_curve", "title": "ROC Curve"})
    if features_df is not None and not features_df.empty:
        sections.extend([
            {"id": "risk_distribution", "title": "위험도 분포"},
            {"id": "risk_count_bar", "title": "위험도별 환자 수"},
        ])
        if "drug_count" in features_df.columns:
            sections.append({"id": "drug_count_distribution", "title": "약물 수 분포"})
        if "yellow_subtype" in features_df.columns:
            sections.append({"id": "yellow_subtype", "title": "Yellow 서브타입"})
        sections.append({"id": "analysis_subject", "title": "분석 대상 정보"})
    elif last_result.get("risk_summary"):
        sections.extend([
            {"id": "risk_distribution", "title": "위험도 분포"},
            {"id": "risk_count_bar", "title": "위험도별 환자 수"},
        ])
    if last_result.get("ddi_means"):
        sections.append({"id": "ddi_severity", "title": "DDI 심각도"})
    elif features_df is not None and not features_df.empty and any(c in features_df.columns for c in ["ddi_contraindicated", "ddi_major", "ddi_moderate", "ddi_minor"]):
        sections.append({"id": "ddi_severity", "title": "DDI 심각도"})
    if metrics.get("classification_report"):
        sections.append({"id": "classification_report", "title": "분류 보고서"})
    if _training_results_rows(training_results):
        sections.append({"id": "sequential_training_results", "title": "이번 순차 학습 결과"})
    if len(_comparison_results(saved_results, training_results)) >= 2:
        sections.append({"id": "model_comparison", "title": "모델 비교"})
    return sections


def build_docx_bytes(last_result: dict,
                     features_df: Optional[pd.DataFrame] = None,
                     saved_results: Optional[list[dict]] = None,
                     training_results: Optional[dict | list[dict]] = None) -> bytes:
    """종합 서비스 보고서 DOCX bytes (4. 결과분석 차트/내용 포함). python-docx 미설치 시 ImportError."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx가 설치되지 않았습니다.")

    doc  = Document()
    now  = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    mname = last_result.get("model_name", "?")
    has_df = features_df is not None and not features_df.empty

    # ── 표지 ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("처방 위험도 예측 종합 보고서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"생성일시: {now}    모델: {mname}")
    doc.add_paragraph()

    # ── 1. 위험도 분포 ────────────────────────────────────────────────────────
    doc.add_heading("1. 위험도 분포", level=1)
    if has_df:
        counts = _count_distribution(features_df)
        total  = len(features_df)
    else:
        risk_summary = last_result.get("risk_summary", {})
        counts = dict(risk_summary)
        total  = sum(counts.values())

    total_d = total or 1

    if has_df and MPL_AVAILABLE:
        assert features_df is not None
        _add_png(doc, _chart_risk_pie(features_df), width_inches=4.5)
        _add_png(doc, _chart_risk_bar_from_counts(features_df["risk_level"].value_counts().to_dict()), width_inches=5.5)
    elif counts and MPL_AVAILABLE:
        _add_png(doc, _chart_pie_from_counts(counts, "위험도 분포 (저장 요약)"), width_inches=4.5)
        _add_png(doc, _chart_risk_bar_from_counts(counts), width_inches=5.5)
    doc.add_paragraph()

    # ── 2. 개입 위계 분포 ─────────────────────────────────────────────────────
    doc.add_heading("2. 개입 위계별 환자 분포", level=1)
    if MPL_AVAILABLE:
        _add_png(doc, _chart_intervention(counts, total_d), width_inches=6.0)
    doc.add_paragraph()

    _add_table(doc,
        ["개입조치", "라벨", "건수", "비율(%)"],
        [[action, lbl, f"{counts.get(lbl.split('·')[0], counts.get(action, 0)):,}",
          f"{counts.get(lbl.split('·')[0], counts.get(action, 0)) / total_d * 100:.2f}"]
         for action, lbl, _ in _INTERVENTION_ROWS]
    )
    doc.add_paragraph()

    # ── 2-1. Yellow 서브타입 세부 분포 ───────────────────────────────────────
    if has_df and MPL_AVAILABLE:
        assert features_df is not None
        _ys_pie = _chart_yellow_pie(features_df)
        _ys_png = _chart_yellow_subtype(features_df)
        _action_png = _chart_action_distribution(features_df)
        if _ys_pie or _ys_png or _action_png:
            doc.add_heading("2-1. Yellow 세분화", level=2)
            if _ys_pie:
                _add_png(doc, _ys_pie, width_inches=5.0)
            if _ys_png:
                _add_png(doc, _ys_png, width_inches=6.0)
            if _action_png:
                doc.add_heading("권장 개입(action) 분포 — Red 포함", level=3)
                _add_png(doc, _action_png, width_inches=6.0)
            if "red_suspect" in features_df.columns:
                rs_count = int(features_df["red_suspect"].eq(True).sum())
                rs_pct = rs_count / len(features_df) * 100 if len(features_df) else 0.0
                doc.add_paragraph(f"Red 의심 (red_suspect=True): {rs_count:,}건 ({rs_pct:.1f}%)")
            doc.add_paragraph()

    # ── 3. DDI 심각도 통계 ────────────────────────────────────────────────────
    ddi_means  = last_result.get("ddi_means") or {}
    if not ddi_means and has_df:
        assert features_df is not None
        ddi_cols = ["ddi_contraindicated", "ddi_major", "ddi_moderate", "ddi_minor"]
        ddi_means = {c: float(features_df[c].mean()) for c in ddi_cols if c in features_df.columns}
    drug_stats = last_result.get("drug_count_stats") or {}
    if not drug_stats and has_df and "drug_count" in features_df.columns:
        dc = features_df["drug_count"].dropna()
        if not dc.empty:
            drug_stats = {"mean": float(dc.mean()), "max": int(dc.max())}

    if ddi_means or drug_stats:
        doc.add_heading("3. DDI 심각도 / 다약제 통계", level=1)
        if ddi_means and MPL_AVAILABLE:
            _add_png(doc, _chart_ddi(ddi_means), width_inches=5.0)
        ddi_label_map = {
            "ddi_contraindicated": "금기 DDI 평균 쌍 수",
            "ddi_major":           "Major DDI 평균 쌍 수",
            "ddi_moderate":        "Moderate DDI 평균 쌍 수",
            "ddi_minor":           "Minor DDI 평균 쌍 수",
        }
        rows = [[lbl, f"{ddi_means[k]:.4f}"]
                for k, lbl in ddi_label_map.items() if k in ddi_means]
        if drug_stats:
            if "mean" in drug_stats:
                rows.append(["약물 수 평균", str(drug_stats["mean"])])
            if "max" in drug_stats:
                rows.append(["약물 수 최대", str(drug_stats["max"])])
        if rows:
            _add_table(doc, ["항목", "값"], rows)
        doc.add_paragraph()

    # ── 4. 피처 중요도 ────────────────────────────────────────────────────────
    fi_data = last_result.get("feature_importance")
    _fi_ok  = fi_data is not None and (
        fi_data.empty is False if hasattr(fi_data, "empty") else bool(fi_data))
    if _fi_ok:
        doc.add_heading("4. 피처 중요도 Top 15", level=1)
        try:
            fi_df  = pd.DataFrame(fi_data) if isinstance(fi_data, list) else fi_data
            fi_top = fi_df.sort_values("importance", ascending=False).head(15)
            if MPL_AVAILABLE:
                _add_png(doc, _chart_feature_importance(fi_df), width_inches=6.0)
            _add_table(doc, ["피처", "한국어명", "중요도"],
                [[str(r["feature"]),
                  _FEAT_LABELS.get(str(r["feature"]), ""),
                  f"{r['importance']:.4f}"]
                 for _, r in fi_top.iterrows()])
        except Exception:
            doc.add_paragraph("피처 중요도 데이터를 표시할 수 없습니다.")
        doc.add_paragraph()

    # ── 5. 모델 성능 지표 ─────────────────────────────────────────────────────
    metrics = last_result.get("metrics", {})
    doc.add_heading("5. 모델 성능 지표", level=1)
    perf_rows = [
        [name, f"{val:.4f}" if isinstance(val, float) else str(val)]
        for name, val in [
            ("F1 (macro)",  metrics.get("f1_macro")),
            ("Accuracy",    metrics.get("accuracy")),
            ("AUC",         metrics.get("roc_auc") or metrics.get("roc_auc_ovr")),
            ("CV 평균",     metrics.get("cv_mean")),
            ("τ_red",       metrics.get("tau_red")),
            ("τ_review",    metrics.get("tau_review")),
        ] if val is not None
    ]
    if perf_rows:
        _add_table(doc, ["지표", "값"], perf_rows)
    train_size = metrics.get("train_size")
    test_size = metrics.get("test_size")
    if train_size is not None or test_size is not None:
        doc.add_paragraph(
            f"학습 {train_size if train_size is not None else '?'}건 | "
            f"테스트 {test_size if test_size is not None else '?'}건"
        )
    doc.add_paragraph()

    training_rows = _training_results_rows(training_results)
    if training_rows:
        doc.add_heading("5-0. 이번 순차 학습 결과", level=2)
        doc.add_paragraph(
            "3단계 모델학습에서 선택된 ML/DL 모델을 순서대로 학습한 결과입니다. "
            "식별자나 원자료 행은 포함하지 않고 모델별 성능 지표만 기록합니다."
        )
        _add_table(
            doc,
            ["순서", "모델", "타겟", "Accuracy", "F1", "AUC", "CV 평균", "학습 수"],
            training_rows,
        )
        doc.add_paragraph()

    if metrics.get("confusion_matrix"):
        doc.add_heading("5-1. 혼동 행렬", level=2)
        cm_val = metrics.get("confusion_matrix")
        if MPL_AVAILABLE:
            _add_png(doc, _chart_confusion_matrix(cm_val, metrics.get("classes")), width_inches=4.8)
        _add_table(doc, ["행", "값"], [[f"Row {i + 1}", str(row)] for i, row in enumerate(cm_val or [])])
        try:
            cm_arr = pd.DataFrame(cm_val).to_numpy()
            if cm_arr.shape == (2, 2):
                tn, fp, fn, tp = cm_arr.ravel()
                precision = tp / (tp + fp) if (tp + fp) > 0 else 0
                recall = tp / (tp + fn) if (tp + fn) > 0 else 0
                f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0
                _add_table(doc, ["지표", "값"], [
                    ["정밀도 (Precision)", f"{precision:.4f}"],
                    ["재현율 (Recall)", f"{recall:.4f}"],
                    ["F1 Score", f"{f1:.4f}"],
                ])
        except Exception:
            pass
        doc.add_paragraph()

        mis_summary = _summarize_misclassification_reasons(metrics, features_df if has_df else None)
        if mis_summary["total"]:
            doc.add_heading("5-1-1. 오판 사유 분석", level=3)
            doc.add_paragraph(
                "개인 식별자와 원자료 식별값은 제외하고, 평가 세트의 안전 feature와 예측 결과만으로 요약했습니다."
            )
            _add_table(doc, ["오판 유형", "건수"],
                       [[k, f"{v:,}건"] for k, v in mis_summary["type_counts"].items()])
            if mis_summary["examples"]:
                _add_table(doc, ["익명번호", "실제", "예측", "확신도", "오판 유형", "추정 사유"], [
                    [e["no"], e["actual"], e["predicted"], e["confidence"], e["type"], e["reason"]]
                    for e in mis_summary["examples"]
                ])
            doc.add_paragraph()

    if metrics.get("cv_scores"):
        cv_scores = list(metrics.get("cv_scores", []))
        doc.add_heading("5-2. 교차검증", level=2)
        if MPL_AVAILABLE:
            _add_png(doc, _chart_cv_scores(cv_scores), width_inches=5.8)
        if cv_scores:
            mean_cv = sum(cv_scores) / len(cv_scores)
            std_cv = (sum((v - mean_cv) ** 2 for v in cv_scores) / len(cv_scores)) ** 0.5
            _add_table(doc, ["지표", "값"], [
                ["평균", f"{mean_cv:.4f}"],
                ["표준편차", f"{std_cv:.4f}"],
                ["최솟값", f"{min(cv_scores):.4f}"],
                ["최댓값", f"{max(cv_scores):.4f}"],
            ])
        doc.add_paragraph()

    roc_data = metrics.get("roc_curve")
    if roc_data and "fpr" in roc_data and "tpr" in roc_data:
        doc.add_heading("5-3. ROC Curve", level=2)
        if MPL_AVAILABLE:
            _add_png(doc, _chart_roc_curve(roc_data, metrics.get("roc_auc")), width_inches=5.0)
        doc.add_paragraph(f"AUC: {metrics.get('roc_auc', metrics.get('roc_auc_ovr', '?'))}")
        doc.add_paragraph()

    report = metrics.get("classification_report")
    if report:
        doc.add_heading("5-4. 분류 보고서", level=2)
        for line in str(report).splitlines():
            doc.add_paragraph(line)
        doc.add_paragraph()

    comparison_results = _comparison_results(saved_results, training_results)
    if len(comparison_results) >= 2:
        doc.add_heading("5-5. 모델 비교", level=2)
        doc.add_paragraph(
            "현재 세션 결과를 우선 포함하고, 저장 이력은 모델별 최신 결과 1건씩만 "
            "뒤에 이어 표시합니다. 같은 모델의 과거 이력은 중복 표시하지 않습니다. "
            "시간은 각 모델의 최신 저장 시각입니다."
        )
        if MPL_AVAILABLE:
            _add_png(doc, _chart_model_comparison(comparison_results), width_inches=6.0)
        rows = []
        for r in comparison_results:
            m = r.get("metrics", {}) or {}
            rows.append([
                str(r.get("timestamp", "?")),
                str(r.get("model_name", "?")),
                str(r.get("model_family", _model_family(r.get("model_name", "?")))),
                str(r.get("target", "?")),
                # 미측정(None) 지표는 §5와 동일하게 "—" 표기 — _metric_float로
                # 감싸면 None→0.0000으로 렌더되어 '최악 성능'처럼 오독됨
                _fmt_report_metric(m.get("accuracy")),
                _fmt_report_metric(m.get("f1_macro")),
                _fmt_report_metric(m.get("roc_auc", m.get("roc_auc_ovr"))),
                _fmt_report_metric(m.get("cv_mean")),
                str(m.get("train_size", 0)),
            ])
        _add_table(doc, ["시간", "모델", "구분", "타겟", "Accuracy", "F1", "AUC", "CV 평균", "학습 수"], rows)
        doc.add_paragraph()

    # ── 6. 분석 메모 ─────────────────────────────────────────────────────────
    doc.add_heading("6. 분석 메모", level=1)
    notes = [
        "본 보고서는 MODE_11_hana 배치 학습 예측 결과 기준입니다.",
        "서빙 실시간 예측과 수치 차이가 있을 수 있습니다.",
    ]
    if not MPL_AVAILABLE:
        notes.append("※ matplotlib 미설치로 차트가 포함되지 않았습니다.")
    for note in notes:
        doc.add_paragraph(note)

    # ── 7. 분석 대상 정보 (마지막 페이지) ─────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("7. 분석 대상 정보", level=1)

    if has_df:
        assert features_df is not None
        analysis_subject_df = _deduplicate_features_by_patient_id(features_df)
        total_n = len(analysis_subject_df)
        rl_dist = analysis_subject_df["risk_level"].value_counts().to_dict() if "risk_level" in analysis_subject_df.columns else {}

        # 카운트 기준 투명화 (③): 피처 행수 vs 고유 환자수. 중복 patient_id 가 있으면
        # §1 위험도분포(행 기준)와 본 표(환자·중복제거 기준)의 총 N 이 달라진다 — 명시 경고.
        raw_n = len(features_df)
        if "patient_id" in features_df.columns:
            _pid = features_df["patient_id"].astype("string").str.strip()
            distinct_pid = int(_pid[_pid != ""].nunique(dropna=True))
        else:
            distinct_pid = raw_n

        target_rows: list[list[str]] = []
        if raw_n != distinct_pid:
            target_rows.append(
                ["⚠ 행/환자 불일치",
                 f"피처 {raw_n:,}행 · 고유 환자 {distinct_pid:,}명 — 중복 patient_id 존재. "
                 "§1 위험도분포는 행 기준, 본 표는 환자(중복제거) 기준."]
            )
        # '총 환자 수' → 다제·중복제거 후임을 라벨에 명시(추출 대상자 N 과의 차이 설명).
        target_rows.append(["분석대상 환자 수 (다제≥5·중복제거 후)", f"{total_n:,}명"])
        for lbl in ["Red", "Yellow", "Green", "Normal"]:
            cnt = rl_dist.get(lbl, 0)
            pct = cnt / total_n * 100 if total_n else 0.0
            target_rows.append([f"  위험도 — {lbl}", f"{cnt:,}명 ({pct:.1f}%)"])

        if total_n > 0:
            if "sex_type" in analysis_subject_df.columns:
                # Raw HANA sex_type contract: 1=남, 2=여. Do not fall back to
                # derived sex_m; other/missing raw values are reported as 미상.
                _sex_vals = analysis_subject_df["sex_type"].astype("string").str.strip()
                male_n = int((_sex_vals == "1").sum())
                female_n = int((_sex_vals == "2").sum())
                unknown_n = total_n - male_n - female_n
                target_rows += [
                    ["성별 — 남", f"{male_n:,}명 ({male_n / total_n * 100:.1f}%)"],
                    ["성별 — 여", f"{female_n:,}명 ({female_n / total_n * 100:.1f}%)"],
                    ["성별 — 미상", f"{unknown_n:,}명 ({unknown_n / total_n * 100:.1f}%)"],
                ]
                if male_n / total_n < 0.20:
                    target_rows.append(
                        ["⚠ 성별 데이터 점검",
                         "남성 비율이 비정상적으로 낮습니다 — 원본 sex_type 값이 누락 또는 비정상 값으로 인해 "
                         "왜곡될 수 있습니다. DEMOGRAPHICS_PATH·sex_type dtype 및 결측 처리 확인."]
                    )
            else:
                target_rows.append(
                    ["⚠ 원본 성별 데이터 없음", "sex_type 컬럼이 없어 성별 집계를 생략했습니다."]
                )

        if "age" in analysis_subject_df.columns:
            age_v = analysis_subject_df["age"].dropna()
            age_v = age_v[age_v >= 0]
            if not age_v.empty:
                target_rows += [
                    ["연령 평균",  f"{age_v.mean():.1f}세"],
                    ["연령 중앙값", f"{age_v.median():.1f}세"],
                    ["연령 범위",  f"{int(age_v.min())}~{int(age_v.max())}세"],
                ]

        if "drug_count" in analysis_subject_df.columns:
            dc = analysis_subject_df["drug_count"].dropna()
            if not dc.empty:
                target_rows += [
                    ["약물 수 평균", f"{dc.mean():.1f}종"],
                    ["약물 수 최대", f"{int(dc.max())}종"],
                ]

        _add_table(doc, ["항목", "값"], target_rows)
        doc.add_paragraph()

        if MPL_AVAILABLE:
            _drug_png = _chart_drug_hist(analysis_subject_df)
            _age_png  = _chart_age_hist(analysis_subject_df)
            if _drug_png:
                _add_png(doc, _drug_png, width_inches=5.5)
            if _age_png:
                _add_png(doc, _age_png, width_inches=5.5)
    else:
        doc.add_paragraph("현재 세션 features_df 없음 — 학습 실행 후 재생성하면 분석 대상 정보가 채워집니다.")

    doc.add_paragraph(
        f"모델: {mname}  |  타겟: {last_result.get('target', '?')}  |  생성일시: {now}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
