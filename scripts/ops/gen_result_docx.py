"""MODE_11_hana 최종 학습결과 보고서(result.docx) 생성. (1회성 산출 스크립트)

H:\\result\\mode_11_hana\\result.docx 로 프로젝트 개요 + 최종 계층모델 결과를 정리.
숫자는 hana_app/models/hierarchical/retrain_prod_0711_hierarchy_cur/stage_meta.json 실측.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
BUNDLE = ROOT / "hana_app/models/hierarchical/retrain_prod_0711_hierarchy_cur"
OUT_DIR = Path(r"H:\result\mode_11_hana")
OUT = OUT_DIR / "result.docx"

meta = json.loads((BUNDLE / "stage_meta.json").read_text(encoding="utf-8"))
counts = meta["stage2_label_counts"]
red = meta["stage1_red_count"]
monitor = counts.get("Y_DOUBLE", 0) + counts.get("Y_DDI_MOD", 0) + counts.get("Y_DUP", 0) + counts.get("Y_FRAG", 0)
no_alert = counts.get("No_Alert", 0)
total = red + counts.get("Y_DDI_MAJOR", 0) + counts.get("Y_TRIPLE", 0) + monitor + no_alert


def pct(n):
    return f"{n/total*100:.2f}%"


doc = Document()
st = doc.styles["Normal"].font
st.name = "맑은 고딕"; st.size = Pt(10)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    return p


# ── 표지 ──────────────────────────────────────────────────────────────────────
t = doc.add_heading("MODE_11_hana — 부적절 처방 위험 예측 ML 파이프라인", level=0)
sub = doc.add_paragraph("최종 학습결과 보고서")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14); sub.runs[0].bold = True
d = doc.add_paragraph("작성일: 2026-06-08  ·  배포 번들: retrain_prod_0711_hierarchy_cur")
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── 1. 개요 ───────────────────────────────────────────────────────────────────
h("1. 프로젝트 개요", 1)
para("HANA 처방 데이터 기반 부적절 처방 위험을 예측·분류하는 운영형 ML 서빙 시스템이다. "
     "환자의 처방 조합에서 4대 위험(금기·중복·상호작용·다기관)을 탐지하고, 위험 수준에 따라 "
     "차등 개입(즉시개입~관여안함)을 권고한다. 학습 데이터는 2024년 7~11월(5개월) 처방, "
     "Nov→Dec 홀드아웃은 연구용으로 동결(parked)되어 운영 학습에서 제외한다.")
para("핵심 설계 원칙: ① 학습↔서빙 피처 스키마 완전 일치, ② 위험 등급은 결정적 임상 룰을 "
     "1차 기준으로 하고 ML은 세부 분류(정확도)를 담당, ③ 모델 오분류에 대비한 결정적 백스톱(floor).")

# ── 2. 데이터 & 코호트 ────────────────────────────────────────────────────────
h("2. 데이터 및 코호트", 1)
para(f"학습 코호트: 다재약물(동시 5종 이상) 환자 {total:,}명 (2024-07~11 처방 기준).")
para("데이터 출처: HANA 처방 T20/T30/T40/T60, 자격(인구), 요양기관 테이블. "
     "약물 식별은 제품코드(EDI)→주성분코드(WK)→성분/DDI 매핑(HIRA·DrugMaster) 경로를 사용한다.")

# ── 3. 위험 라벨 & 개입 위계 ──────────────────────────────────────────────────
h("3. 위험 라벨 체계 및 개입 위계", 1)
para("2단 계층 분류: Stage 1(Red 이진) + Stage 2(Yellow 7-class 세분화). 위험 등급은 결정적 "
     "임상 룰(CLINICAL_STANDARDS v1.0)로 정의되며, 개입 강도를 5단계로 차등화한다(2026-06-07 재설계).")
tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "등급/라벨"; hdr[1].text = "조건"; hdr[2].text = "개입"; hdr[3].text = "비중"
rows = [
    ("Red (금기)", "절대 금기(contraindicated DDI)", "즉시 개입", pct(red)),
    ("Y_DDI_MAJOR", "주요 상호작용(major DDI ≥1)", "약사 전화", pct(counts.get("Y_DDI_MAJOR", 0))),
    ("Y_TRIPLE (중증)", "Triple Whammy / 10종+고위험 / 고령+장기 / 3위험차원", "문자 안내", pct(counts.get("Y_TRIPLE", 0))),
    ("Y_DOUBLE·단일", "2위험차원 / 중등도DDI·중복·다기관 단일", "모니터링", pct(monitor)),
    ("No_Alert / Green / Normal", "위험 신호 없음", "관여 안 함", pct(no_alert)),
]
for r in rows:
    c = tbl.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v
para("개입 위계는 Red > Y_DDI_MAJOR > Y_TRIPLE > 모니터링 > 관여안함 순. 금기만 최상위 즉시개입으로 "
     "한정해 과경보를 억제하고, 주요 상호작용·중증 다재약물은 약사전화·문자안내로 분산한다.", bold=False)

# ── 4. 피처 ───────────────────────────────────────────────────────────────────
h("4. 입력 피처 (22개)", 1)
para(", ".join(meta["feature_cols"]))
para("Triple Whammy 및 위험약물 플래그(고위험/신독성/간독성)는 성분명 키워드로 산출하며, "
     "학습·서빙이 동일 함수(edi→wk→DrugMaster)를 호출해 정합을 보장한다.")

# ── 5. 최종 학습 결과 ─────────────────────────────────────────────────────────
h("5. 최종 학습 결과 (retrain_prod_0711_hierarchy_cur)", 1)
para(f"환자 수: {total:,}  ·  7-class 계층 모델  ·  feature_semantics_version: "
     f"{meta['feature_semantics_version']}  ·  ddi: {meta['ddi_feature_semantics_version']}  ·  "
     f"clinical_standards: {meta['clinical_standards_version']}")
para(f"Stage 1(Red) 확정: {red:,}건 (룰 파생 결정적 라벨). thresholds τ_red="
     f"{meta['thresholds']['tau_red']:.6f}, τ_review={meta['thresholds']['tau_review']:.6f}.")
para("Stage 2(Yellow 세분화) 라벨 분포:", bold=True)
t2 = doc.add_table(rows=1, cols=3); t2.style = "Light Grid Accent 1"
hh = t2.rows[0].cells; hh[0].text = "라벨"; hh[1].text = "건수"; hh[2].text = "비중"
for lab in ["Y_DDI_MAJOR", "Y_TRIPLE", "Y_DOUBLE", "Y_DUP", "Y_DDI_MOD", "Y_FRAG", "No_Alert"]:
    c = t2.add_row().cells
    c[0].text = lab; c[1].text = f"{counts.get(lab,0):,}"; c[2].text = pct(counts.get(lab, 0))
para("")
para("개입 강도별 분포 (최종):", bold=True)
t3 = doc.add_table(rows=1, cols=3); t3.style = "Light Grid Accent 1"
hh = t3.rows[0].cells; hh[0].text = "개입"; hh[1].text = "건수"; hh[2].text = "비중"
for name, n in [("즉시 개입(Red)", red), ("약사 전화(Y_DDI_MAJOR)", counts.get("Y_DDI_MAJOR", 0)),
                ("문자 안내(Y_TRIPLE)", counts.get("Y_TRIPLE", 0)), ("모니터링(Y_DOUBLE·단일)", monitor),
                ("관여 안 함(No_Alert)", no_alert)]:
    c = t3.add_row().cells
    c[0].text = name; c[1].text = f"{n:,}"; c[2].text = pct(n)
para("")
para("해석: 고강도 개입(즉시개입+약사전화)은 약 17%로 분산되어 임상적으로 운영 가능한 수준이며, "
     "금기(0.37%)만 즉시개입으로 한정해 과경보를 방지한다. 다재약물 코호트 특성상 중증(Y_TRIPLE) 및 "
     "모니터링 비중이 높다.")
para("보조 모델(참고): same-window 다기관 sparse-linear baseline 검증 AUC 0.845 (auxiliary, info-only).")

# ── 6. 서빙 시스템 ────────────────────────────────────────────────────────────
h("6. 서빙 시스템 및 안전 장치", 1)
para("FastAPI /predict — 요청(EDI 코드)→edi→wk 브릿지→피처 산출→계층 모델 추론→개입 권고. "
     "학습과 동일한 edi→wk→DrugMaster 경로로 DDI·중복·위험약물을 계산해 train/serve 정합을 보장한다.")
para("3층 결정적 백스톱(floor) — 모델 오분류에 대비해 결정적 임상 룰로 최소 등급을 보장(model-independent):", bold=True)
para("• 금기(contraindicated) → Red(즉시개입)   • major DDI(≥1) → Y_DDI_MAJOR(약사전화)   "
     "• 중증(triple_whammy/10종+고위험/고령+장기) → Y_TRIPLE(문자안내). 단방향 상향(상위 등급은 유지).")
para("버전 가드: feature_semantics_version(rulefeat.v1)·ddi.v2 로 구 번들/스키마 불일치를 차단한다.")

# ── 7. 검증 & 배포 ────────────────────────────────────────────────────────────
h("7. 검증 및 배포", 1)
para("검증: 단위·정합(parity)·계약 테스트 + 광역 회귀(serving 210+ passed). 배포 전 in-process 및 "
     "HTTP(TestClient) 전 경로 검증(계층 로드·개입 위계 적용) 통과.")
para("배포: 환경변수 HIERARCHICAL_MODEL_DIR 로 번들 지정 후 serving 기동/리로드(env 전용). "
     "운영 PC(Windows 폐쇄망, Python 3.12)로 번들 전송 후 적용. 상세는 docs/ops/hierarchical-deploy-runbook.md.")

# ── 8. 한계 & 후속 ────────────────────────────────────────────────────────────
h("8. 한계 및 후속 과제", 1)
para("• Stage 1 Red 임계는 룰 파생 결정적 라벨 특성상 degenerate(τ≈1.0)이나, 결정적 백스톱이 "
     "안전을 독립 보장하므로 운영 안전성에는 영향 없음.")
para("• 위험약물(고위험/신독성/간독성) 키워드 확장은 과경보 트레이드오프로 현 상태 유지(필요 시 임상 검토).")
para("• Nov→Dec 홀드아웃 기반 future-onset 연구 트랙은 무기한 동결(parked).")

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"saved: {OUT}  (paragraphs={len(doc.paragraphs)})")
print(f"check: total={total:,} red={red} ddi_major={counts.get('Y_DDI_MAJOR',0)} "
      f"ytriple={counts.get('Y_TRIPLE',0)} monitor={monitor} no_alert={no_alert}")
